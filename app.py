from flask import Flask, render_template, request, jsonify, redirect, url_for, render_template_string, session,send_from_directory, flash, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, join_room, leave_room, send
from PIL import Image
import mysql.connector
import base64
from inference import get_model
import supervision as sv

import zipfile
import os
from geopy.geocoders import Nominatim


import random

import datetime

import requests

import matplotlib
matplotlib.use('agg')
import matplotlib.pyplot as plt
import string
from io import BytesIO

from docx import Document
from docx.shared import Inches
import winreg as reg
import pandas as pd

from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows

API_KEY = "5c9af440b34248e955e93d77520f9410"


lat = "15.3627"
lon = "120.8838"
BASE_URL = "https://api.openweathermap.org/data/2.5/weather?lat="+ lat + "&lon="+ lon +"&exclude=hourly,daily&appid="+ API_KEY


response = requests.get(BASE_URL).json()


connection = mysql.connector.connect(
    host='localhost',
    port='3306',
    database='new_caps',
    user='root',
    password='',
    connection_timeout=28800
    
)

cursor = connection.cursor(buffered=True)

app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = 'secret!'

socketio = SocketIO(app, cors_allowed_origins="*")

app.config['UPLOAD_FOLDER'] = r'C:\xampp\htdocs\backup capstone\xammp_projects\static\upload_folder'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 




def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_random_code(length=6):
    letters_and_digits = string.ascii_uppercase + string.digits
    return ''.join(random.choice(letters_and_digits) for i in range(length))





def kelvin_to_cs_fr(kelvin):
    cs = kelvin - 273.15

    fr = cs *(9/5) + 32
    return cs, fr


@app.route('/register_un', methods=['GET', 'POST'])
def username():
    text = ''

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        cursor.execute('''SELECT * FROM tbl_user WHERE username = %s AND password = %s''', (username, password,))
        check= cursor.fetchall()
        cursor.execute('''SELECT * FROM tbl_user WHERE username = %s''', (username,))
        check_name= cursor.fetchall()
        if ' ' in username:
            text = 'Username must not contain spaces.'
        elif ' ' in password:
            text = 'Password must not contain spaces.'
        elif check:
            text = 'This account is already registered.'
        elif check_name:
            text = 'This username is already registered. Try another username.'
        elif len(username) < 4:
            text = 'The characters in username must be 4 to 20.'
       
        elif len(password) < 4:
            text = 'The characters in password must be 4 to 20.'
        
        else:
            # 7 is default - nagreregister palang 
            role_id = 7
            current_timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('INSERT INTO tbl_user (username,  role_id, created_at, password) VALUES (%s, %s, %s, %s)', (username, role_id, current_timestamp, password))
            connection.commit()
            cursor.execute('SELECT user_id FROM tbl_user ORDER BY user_id DESC LIMIT 1')
            user_id = cursor.fetchone()[0]
            return redirect(url_for('farmback', user_id = user_id))

    return render_template('register_un.html', text =text)

@app.route('/farmback', methods=['GET', 'POST'])
def farmback():

    user_id = request.args.get('user_id')
    print(user_id)
    if request.method == 'POST':
        role_id = request.form['role_id']
        user_id = request.form['user_id']
        cursor.execute('UPDATE tbl_user SET role_id = %s WHERE user_id = %s', (role_id, user_id))
        connection.commit()
        if role_id == '7':
            print(user_id)
            return redirect(url_for("register_farmer", user_id = user_id))
        elif role_id == '1':
            return redirect(url_for("register_backyard", user_id = user_id))
        else: 
            text = 'ERROR'
            return render_template("farmback.html", text=text)
    return render_template("farmback.html", user_id = user_id)



@app.route('/register_farmer', methods=['GET', 'POST'])
def register_farmer():
    user_id = request.args.get('user_id')
    print(user_id)
    if request.method == 'POST':
        user_id = request.form['user_id']
        
        fname = request.form['fname']
        mname = request.form['mname']
        lname = request.form['lname']
        province = request.form['province-dropdown']
        city = request.form['city-dropdown']
        brgy = request.form['barangay-dropdown']
        strt_add = request.form['strt_add']
        cont_num = request.form['cont_num']

        cursor.execute('SELECT * FROM user_details WHERE fname = %s AND mname = %s AND lname = %s and province =%s AND city = %s AND brgy= %s AND strt_add = %s AND cont_num =%s'
        ,(fname, mname, lname, province, city, brgy, strt_add, cont_num,))
        results = cursor.fetchall()

        if results:
            text = 'You already have an account'
            return render_template('register_farmer.html', user_id = user_id, text = text)

        else:
            cursor.execute('INSERT INTO user_details (fname, mname, lname, province, city, brgy, strt_add, cont_num, user_id) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)', 
                        (fname, mname, lname, province, city, brgy, strt_add, cont_num, user_id))
            connection.commit()
            return redirect(url_for('ask', user_id = user_id))
    return render_template('register_farmer.html', user_id = user_id)


@app.route('/register_backyard', methods=['GET', 'POST'])
def register_backyard():
    user_id = request.args.get('user_id')

    if request.method == 'POST':
        user_id = request.form['user_id']
        fname = request.form['fname']
        mname = request.form['mname']
        lname = request.form['lname']
        province = request.form['province-dropdown']
        city = request.form['city-dropdown']
        brgy = request.form['barangay-dropdown']
        strt_add = request.form['strt_add']
        cont_num = request.form['cont_num']
        country = "Philippines"

        address = f"{brgy}, {city}, {province}, {country}"
        print(f"Geocoding address: {address}")
        cursor.execute('SELECT * FROM user_details WHERE fname = %s AND mname = %s AND lname = %s and province =%s AND city = %s AND brgy= %s AND strt_add = %s AND cont_num =%s'
        ,(fname, mname, lname, province, city, brgy, strt_add, cont_num,))
        results = cursor.fetchall()
        if results:
            text = 'You already have an account'
            return render_template('register_farmer.html', user_id = user_id, text = text)


        else:
            cursor.execute("""INSERT INTO user_details (fname, mname, lname, province, city, brgy, strt_add, cont_num, user_id) 
            VALUES ( %s, %s, %s, %s, %s, %s,%s, %s, %s)""", (fname, mname, lname, province, city, brgy, strt_add, cont_num, user_id,))
            connection.commit()

        return render_template('login.html')
        

       
    return render_template('register_backyard.html', user_id=user_id)



@app.route('/map', methods=['GET'])
def map():

    user_id = session['id']


    if check_person_in_ud_farm(user_id):
        valid_rooms = 1
         
    else:
        valid_rooms = 0

         
    
    query1 = '''
        SELECT f.farm_name, f.latitude, f.longitude
        FROM farm f
    '''
    cursor.execute(query1)
    results = cursor.fetchall()

    map_data = []
    for row in results:
        map_data.append({
            'farm_name': row[0],
            'latitude': row[1],
            'longitude': row[2],
        })
    

    return render_template('map.html', coords=map_data, valid_rooms = valid_rooms)




@app.route('/search_farms', methods=['POST'])
def search_farms():
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')

    search_radius = 1000

    query2 = '''
        SELECT f.farm_name, f.latitude, f.longitude, hs.status_name, COUNT(d.detection_id) AS disease_count
        FROM farm f
        LEFT JOIN detection d ON d.farm_id = f.farm_id
        LEFT JOIN health_status hs ON d.health_status_id = hs.health_status_id
        WHERE ST_Distance_Sphere(POINT(f.longitude, f.latitude), POINT(%s, %s)) <= %s
        AND d.health_status_id IS NOT NULL
        AND d.health_status_id != 8 
        AND d.health_status_id != 9
        ORDER BY disease_count DESC
    '''

    cursor.execute(query2, (longitude, latitude, search_radius))
    results = cursor.fetchall()
    print(results)
    farm_data = []

    if results:

       for row in results:
            farm_data.append({
                'farm_name': row[0],
                'latitude': row[1],
                'longitude': row[2],
                'health_status_name': row[3],
                'disease_count': row[4]
            })
    else:
        farm_data = []


    print(farm_data)

    return jsonify(farm_data)










@app.route('/ask', methods=['GET', 'POST'])
def ask():
    user_id = request.args.get('user_id')
    print('this is user id in ask', user_id)

    if request.method == 'POST':
        user_id = request.form['user_id']

        ask = int(request.form['ask'])
        print(ask)
        if ask == 1:
            # 6 farm owner sya
            cursor.execute('UPDATE tbl_user SET role_id = 6 WHERE user_id = %s', (user_id,))
            connection.commit()
           
            return redirect(url_for('register_farm', user_id = user_id))
        elif ask == 2:
            # 5 farmer lang
            cursor.execute('UPDATE tbl_user SET role_id = 5 WHERE user_id = %s', (user_id,))
            connection.commit()
            return redirect(url_for('nofarm', user_id = user_id))
    return render_template('ask.html', user_id = user_id)




@app.route('/nofarm', methods=['GET', 'POST'])
def nofarm():
    user_id = request.args.get('user_id') 
    farm_id = request.args.get('farm_id') 

    if request.method == 'POST':
        farm_id = request.form['farm_id'] 
        user_id = request.form.get('user_id', user_id) 
        print('userid', user_id)

        if not user_id:  
            user_id = session.get('id')  
            if user_id:
                print('User ID is in session:', user_id)
            else:
                print('User ID not found in session')

        return redirect(url_for('learnmore', user_id=user_id, farm_id=farm_id)) 

    cursor.execute('SELECT * FROM farm WHERE farm_status = "active"')
    farm = cursor.fetchall()

    farm_with_images = []
    for f in farm:
        if f[9]:  
            new_fpath = f[9].replace('/uploads/', '')
        else:
            new_fpath = None 

        farm_with_images.append(f + (new_fpath,))  

    return render_template('nofarm.html', farm=farm_with_images, user_id=user_id, farm_id=farm_id)

    


@app.route('/learnmore', methods=['GET', 'POST'])
def learnmore():

    farm_id = request.args.get('farm_id')
    user_id = request.args.get('user_id')

    print(f'Initial farm_id from GET: {farm_id}')
    print(f'Initial user_id from GET: {user_id}')
    cursor.execute('''SELECT farm_id, farm_name, province, city, brgy, strt_add, zip_code
                      FROM farm
                      WHERE farm_id = %s LIMIT 1''', (farm_id,))
    farm = cursor.fetchone()
    print('Farm data in learnmore:', farm)
    
    cursor.execute('''SELECT ud.fname, ud.mname, ud.lname FROM user_details ud 
                      INNER JOIN ud_farm uf on uf.user_details_id = ud.user_details_id
                      WHERE uf.farm_id = %s AND uf.status = "Owner" LIMIT 1''', (farm_id,))
    owner = cursor.fetchone()

    if owner:
        fname, mname, lname = owner
    else:
        fname, mname, lname = None, None, None 
    

    if request.method == 'POST':
        farm_id = request.form.get('farm_id', farm_id)
        user_id = request.form.get('user_id', user_id)
        print(f'Updated farm_id from POST: {farm_id}')
        print(f'Updated user_id from POST: {user_id}')

        if 'id' in session: 
            user_id = session.get('id') 
            if user_id:
                print('User ID is in session:', user_id)
            else:
                print('User ID not found in session')


        if 'id' not in session:
            cursor.execute('UPDATE tbl_user SET role_id = 5 WHERE user_id = %s', (user_id,))
            connection.commit()
            print(f'Role updated for user_id: {user_id}')
        else:
            print('Session exists, skipping role update')

        cursor.execute('SELECT user_details_id FROM user_details WHERE user_id = %s', (user_id,))
        result = cursor.fetchone()
        user_details_id = result[0] if result else None


        cursor.execute('''SELECT * FROM ud_farm WHERE user_details_id = %s AND farm_id = %s''', (user_details_id,farm_id,))
        check_ud = cursor.fetchall()

        if check_ud:
            text = 'You are already in the farm!'
            return render_template('learnmore.html', farm=farm, user_id=user_id, farm_id=farm_id, text = text)




        if user_details_id and farm_id:
            cursor.execute('SELECT uf.group_id FROM ud_farm uf WHERE uf.farm_id = %s', (farm_id,))
            result = cursor.fetchone()
            group_id = result[0] if result else None

            cursor.execute('INSERT INTO ud_farm (user_details_id, farm_id, status, group_id) VALUES (%s, %s, "member pending", %s)', 
                           (user_details_id, farm_id, group_id))
            connection.commit()

        return redirect(url_for('login'))



    return render_template('learnmore.html', farm=farm, user_id=user_id, farm_id=farm_id, fname = fname, mname = mname, lname = lname)

        



@app.route('/register_farm', methods=['GET', 'POST'])
def register_farm():
    user_id = request.args.get('user_id')


    print(user_id)

    text = ''
    if request.method == 'POST': 
        print("Form data keys:", request.form.keys())
        user_id = request.form['user_id']
        print('user id sa POST', user_id)

        farm_name = request.form['farm_name']
        province = request.form['province-dropdown']
        city = request.form['city-dropdown']
        brgy = request.form['barangay-dropdown']
        strt_add = request.form['strt_add']
        zip_code = request.form['zip-code']
        f_img = request.files['f_img']

        cursor.execute('SELECT user_details_id FROM user_details WHERE user_id = %s', (user_id,))
        user_details = cursor.fetchone()
        user_details_id = user_details[0]


        if f_img:
            filename = f_img.filename
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            f_img.save(file_path)
            image_url = url_for('uploaded_file', filename=filename)

        else:
            image_url = None



        country = 'Philippines'
        address = f"{brgy}, {city}, {province}, {country}"
        print(f"Geocoding address: {address}")

        geolocator = Nominatim(user_agent="JOSHAPP")
        location = geolocator.geocode(address)

        if location:
            latitude = location.latitude
            longitude = location.longitude
            cursor.execute('INSERT INTO farm (farm_name, province, city, brgy, strt_add, zip_code, latitude, longitude, f_img, farm_status) VALUES (%s, %s, %s, %s, %s, %s, %s,%s, %s, "active")', 
                       (farm_name, province, city, brgy, strt_add, zip_code, latitude, longitude, image_url))
            connection.commit()
            cursor.execute('SELECT farm_id FROM farm ORDER BY farm_id DESC LIMIT 1')
            farm_id = cursor.fetchone()[0]



            collab_group_code = generate_random_code()

            cursor.execute("INSERT INTO tbl_group (group_code) values (%s)",(collab_group_code,))
            connection.commit()
            cursor.execute("SELECT group_id FROM tbl_group ORDER BY group_id DESC LIMIT 1")
            owner_group_id = cursor.fetchone()[0]

            cursor.execute('INSERT INTO ud_farm VALUES(NULL, %s, %s, "Owner", %s)', (user_details_id, farm_id, owner_group_id))
            connection.commit()
            print('THIS IS GROUP ID', owner_group_id)

            return redirect(url_for('login'))
            

        else:
            print("No geocoding results found.")
            text = 'Please Make sure the address is correct'


        
    return render_template('register_farm.html', user_id = user_id, text = text)



@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'id' in session:
        sid = session['id']
        if 'farm_id' in session:
            farm_id = session['farm_id']
            if farm_id is not None:  
                cursor.execute('''SELECT uf.status 
                                    FROM tbl_user u
                                    INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                    INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                    WHERE u.user_id = %s AND uf.farm_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                                (sid, farm_id,))  

                result = cursor.fetchone()  
                status = result[0] if result else None  
                print("This is the status:", status)  
            else:
                cursor.execute('''SELECT uf.status 
                                    FROM tbl_user u
                                    INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                    INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                    WHERE u.user_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                                (sid,))  

                result = cursor.fetchone()  
                status = result[0] if result else None  
                print("This is the status:", status) 
        else:
            status = None
            
        role_id = session['role_id']
        if check_person_in_ud_farm(sid):
            valid_rooms = 1
        else:
            valid_rooms = 0
        
        cursor.execute('SELECT u.username, ud.fname, ud.mname, ud.lname, ud.province, ud.city, ud.brgy, ud.strt_add, ud.profile_pic ,u.password FROM tbl_user u INNER JOIN user_details ud on ud.user_id = u.user_id WHERE u.user_id = %s', (sid,))
        info = cursor.fetchone()

        cursor.execute('''SELECT tr.role_name FROM tbl_role tr INNER JOIN tbl_user u on u.role_id = tr.role_id WHERE u.user_id = %s''', (sid,))
        role_name = cursor.fetchone()[0]

        farm_info = []
        if role_name in ['farm_owner', 'user_farmer']:
            cursor.execute('''SELECT f.farm_name, f.province, f.city, f.brgy, f.strt_add FROM farm f INNER JOIN ud_farm uf
                               ON uf.farm_id = f.farm_id INNER JOIN user_details ud ON ud.user_details_id = uf.user_details_id
                               WHERE ud.user_id = %s''', (sid,))
            farm_info = cursor.fetchall()

        if request.method == 'POST':
            file = request.files.get('file')
            if file and allowed_file(file.filename):
                filename = file.filename
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                cursor.execute('UPDATE user_details SET profile_pic = %s WHERE user_id = %s', (filename, sid))
            
            username = request.form.get('username')
            password = request.form.get('password')
            fname = request.form.get('fname')
            mname = request.form.get('mname')
            lname = request.form.get('lname')
            province = request.form.get('province')
            city = request.form.get('city')
            brgy = request.form.get('brgy')
            strt_add = request.form.get('strt_add')
            
            cursor.execute('UPDATE tbl_user SET username = %s WHERE user_id = %s', (username, sid))
            cursor.execute('UPDATE tbl_user SET password = %s WHERE user_id = %s', (password, sid))
            cursor.execute('''UPDATE user_details SET fname = %s, mname = %s, lname = %s, province = %s, 
                              city = %s, brgy = %s, strt_add = %s WHERE user_id = %s''',
                           (fname, mname, lname, province, city, brgy, strt_add, sid))
            connection.commit()
            return redirect(url_for('profile'))


        
        
        return render_template('profile.html', info=info, farm_info=farm_info, user_type=role_name, valid_rooms=valid_rooms, status = status)
    
    return redirect(url_for('login'))






def is_valid_xlsx(file_path):
    """Check if the file is a valid .xlsx (zip) file."""
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.testzip()
        return True
    except zipfile.BadZipFile:
        return False


backup_dir = r'C:\xampp\htdocs\backup capstone\xammp_projects\static\data_backup'

if not os.path.exists(backup_dir):
    os.makedirs(backup_dir)

def save_to_excel(data, sheet_name):
    file_path = os.path.join(backup_dir, f'{sheet_name}_backup.xlsx')

    if os.path.exists(file_path):
        wb = load_workbook(file_path)
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.create_sheet(sheet_name)
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

    for r in dataframe_to_rows(data, index=True, header=True):
        ws.append(r)

    wb.save(file_path)
    print(f"Data saved to {file_path}")

@app.route('/delete_account', methods=['POST'])
def delete_account():
    if 'id' in session:
        sid = session['id']
        role_id = session['role_id']

        cursor.execute('''SELECT uf.user_details_id, uf.farm_id, uf.status 
                          FROM ud_farm uf 
                          INNER JOIN user_details ud on uf.user_details_id = ud.user_details_id
                          INNER JOIN tbl_user u on u.user_id = ud.user_id 
                          WHERE u.user_id = %s 
                          ORDER BY uf.user_details_id DESC LIMIT 1''', (sid,))
        res = cursor.fetchone()

        user_data = {}

        cursor.execute('SELECT * FROM tbl_user WHERE user_id = %s', (sid,))
        user_data['tbl_user'] = pd.DataFrame(cursor.fetchall())
        cursor.execute('SELECT * FROM user_details WHERE user_id = %s', (sid,))
        user_data['UserDetails'] = pd.DataFrame(cursor.fetchall())
        cursor.execute('''SELECT * FROM detection WHERE user_id = %s''', (sid,))
        user_data['Detection'] = pd.DataFrame(cursor.fetchall())
        
        if res:
            check_farm = res[0]
            status = res[2]
            farm_id = res[1]
            print(check_farm)

            cursor.execute('''SELECT * FROM tbl_message WHERE ud_farm_id IN (
                                SELECT ud_farm_id FROM ud_farm WHERE user_details_id = %s)''', (check_farm,))
            user_data['Messages'] = pd.DataFrame(cursor.fetchall())

            cursor.execute('''SELECT * FROM history WHERE detection_id IN (
                                SELECT detection_id FROM detection WHERE user_id = %s)''', (sid,))
            user_data['History'] = pd.DataFrame(cursor.fetchall())

            for sheet_name, data in user_data.items():
                save_to_excel(data, sheet_name)

            if check_farm:
                cursor.execute('''DELETE m FROM tbl_message m INNER JOIN ud_farm uf ON m.ud_farm_id = uf.ud_farm_id
                                  WHERE uf.user_details_id = %s''', (check_farm,))
                connection.commit()

                if status == 'Owner' and role_id in [6]:
                    cursor.execute('''UPDATE farm f
                                      INNER JOIN ud_farm uf ON uf.farm_id = f.farm_id
                                      INNER JOIN user_details ud ON uf.user_details_id = ud.user_details_id
                                      SET f.farm_status = %s
                                      WHERE uf.user_details_id = %s AND uf.status = %s''', 
                                   ('Pending For Deletion', check_farm, 'Owner'))
                    connection.commit()

                cursor.execute('DELETE FROM ud_farm WHERE user_details_id = %s', (check_farm,))
                connection.commit()

        else:
            for sheet_name, data in user_data.items():
                save_to_excel(data, sheet_name)

        cursor.execute('DELETE FROM history WHERE detection_id IN (SELECT detection_id FROM detection WHERE user_id = %s)', (sid,))
        connection.commit()

        cursor.execute('DELETE FROM user_details WHERE user_id = %s', (sid,))
        connection.commit()

        cursor.execute('DELETE FROM detection WHERE user_id = %s', (sid,))
        connection.commit()

        cursor.execute('DELETE FROM tbl_user WHERE user_id = %s', (sid,))
        connection.commit()

        session.clear()

        return redirect(url_for('login'))

    return redirect(url_for('login'))

@app.route('/members', methods=['GET', 'POST'])
def members():
    if 'id' in session:
        sid = session['id']
        farm_id = session['farm_id']

        if farm_id is not None:
            cursor.execute('''SELECT uf.status 
                               FROM tbl_user u
                               INNER JOIN user_details ud ON u.user_id = ud.user_id 
                               INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                               WHERE u.user_id = %s AND uf.farm_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                           (sid, farm_id))

            result = cursor.fetchone()
            status = result[0] if result else None
            print("This is the status:", status)
        else:
            cursor.execute('''SELECT uf.status 
                               FROM tbl_user u
                               INNER JOIN user_details ud ON u.user_id = ud.user_id 
                               INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                               WHERE u.user_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                           (sid,))

            result = cursor.fetchone()
            status = result[0] if result else None
            print("This is the status:", status)

        cursor.execute('''SELECT tr.role_name 
                          FROM tbl_role tr 
                          INNER JOIN tbl_user u on u.role_id = tr.role_id 
                          WHERE u.user_id = %s''', (sid,))
        role_name = cursor.fetchone()[0]
        
        if check_person_in_ud_farm(sid):
            valid_rooms = 1
        else:
            valid_rooms = 0

        cursor.execute('''SELECT uf.farm_id, f.farm_name, uf.status
                          FROM tbl_user u 
                          INNER JOIN user_details ud ON u.user_id = ud.user_id 
                          INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                          INNER JOIN farm f ON f.farm_id = uf.farm_id
                          WHERE u.user_id = %s''', (sid,))
        farms = cursor.fetchall()
        print('Farms associated with user: ', farms)

        valid_farms = []
        for farm in farms:
            farm_id = farm[0]
            farm_name = farm[1]
            farm_status = farm[2]

            if farm_status not in ['rejected', 'member pending']:
                valid_farms.append({'farm_id': farm_id, 'farm_name': farm_name})

        all_members_by_farm = []

        for farm in valid_farms:
            farm_id = farm['farm_id']
            farm_name = farm['farm_name']
            print('Processing farm id: ', farm_id)

            cursor.execute('''SELECT u.username, ud.fname, ud.mname, ud.lname, 
                                      f.farm_name, uf.status, uf.group_id, uf.user_details_id
                              FROM tbl_user u 
                              INNER JOIN user_details ud ON u.user_id = ud.user_id 
                              INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                              INNER JOIN farm f ON f.farm_id = uf.farm_id
                              WHERE uf.farm_id = %s''', (farm_id,))
            members = cursor.fetchall()

            if members:
                all_members_by_farm.append({
                    'farm_name': farm_name,
                    'members': members
                })

        group_id = None
        print('This is all_members_by_farm:', all_members_by_farm)

        if request.method == 'POST':
            user_details_id = request.form['user_details_id']
            permission = request.form['permission']

            for member in all_members_by_farm:
                for m in member['members']:
                    if str(m[7]) == user_details_id:
                        group_id = m[6]
                        break

            print('POST: group id:', group_id, 'permission:', permission, 'user_details_id:', user_details_id)

            if group_id is None:
                print("Error: Group ID is None, possible data mismatch.")

            if permission == 'accept' and group_id:
                cursor.execute('''UPDATE ud_farm SET status = "member", group_id = %s WHERE user_details_id = %s''', 
                               (group_id, user_details_id,))
                connection.commit()

            elif permission == 'reject':
                cursor.execute('UPDATE ud_farm SET status = "rejected" WHERE user_details_id = %s', (user_details_id,))
                connection.commit()

        return render_template('members.html', 
                               all_members_by_farm=all_members_by_farm, 
                               status=status, 
                               group_id=group_id, 
                               user_type=role_name,
                               valid_rooms=valid_rooms)



@app.route('/admin/delete_farm', methods=['POST'])
def delete_farm():
    if 'admin' in session:
        farm_id = request.form['farm_id']
        farm_status = request.form['farm_status']

        cursor.execute('SELECT * FROM ud_farm WHERE farm_id = %s', (farm_id,))
        farm_data = cursor.fetchall()

        if farm_data:
            flash(f'Farm ID {farm_id} has existing farmers.', 'success')
            return redirect(url_for('update_farm_status'))
                        
        else:
            columns = [desc[0] for desc in cursor.description]
            farm_df = pd.DataFrame(farm_data, columns=columns)

            save_to_excel(farm_df, "Deleted Farms")
            cursor.execute('DELETE FROM farm WHERE farm_id = %s', (farm_id,))
            connection.commit()


        flash(f'Farm ID {farm_id} status updated to {farm_status}.', 'success')
        return redirect(url_for('update_farm_status'))
    else:
        flash('You must be an admin to access this page.', 'danger')
        return redirect(url_for('login'))
       

@app.route('/admin/update_farm_status', methods=['GET', 'POST'])
def update_farm_status():
    if 'id' in session:
        if request.method == 'POST':
            farm_id = request.form['farm_id']
            farm_status = request.form['farm_status']

            cursor.execute('''UPDATE farm 
                              SET farm_status = %s 
                              WHERE farm_id = %s''', (farm_status, farm_id))
            connection.commit()
            
            flash(f'Farm ID {farm_id} status updated to {farm_status}.', 'success')
            return redirect(url_for('update_farm_status'))

        cursor.execute('''SELECT farm_id, farm_name, farm_status, f_img FROM farm''')
        farms = cursor.fetchall()

        farm_with_images = []
        for farm in farms:
            farm_id = farm[0]
            farm_name = farm[1]
            farm_status = farm[2]
            farm_image = farm[3] 

            if farm_image:
                farm_image = farm_image.replace('/uploads/', '')
                image_url = farm_image 
            else:
                image_url = None

            farm_with_images.append({
                'farm_id': farm_id,
                'farm_name': farm_name,
                'farm_status': farm_status,
                'image_url': image_url
            })
        return render_template('admin_farm_status.html', farms=farm_with_images)
    
    else:
        flash('You must be an admin to access this page.', 'danger')
        return redirect(url_for('login'))






@app.route('/contact', methods=['GET', 'POST'])
def contact():

    return render_template('contact.html')


@app.route('/about', methods=['GET', 'POST'])
def about():

    return render_template('about.html')




@app.route('/admin_test', methods=['GET', 'POST'])
def admin_test():
    if 'id' in session:
        sid = session['id']
        cursor.execute("SELECT username FROM tbl_user WHERE user_id = %s", (sid,))
        result = cursor.fetchone()

        image_url = None
        formatted_labels = None

        if result:
            username = result[0]
        else:
            return redirect(url_for('login'))

        label = ''
        if request.method == 'POST':
            file = request.files['file']
            if file and allowed_file(file.filename):
                filename = file.filename
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                image = Image.open(file_path)

                model = get_model(model_id="calamansi-0bxbw/3")
                results = model.infer(image)

                detections = sv.Detections.from_inference(results[0].dict(by_alias=True, exclude_none=True))

                labels = [f"{class_name} {confidence:.2f}" for class_name, confidence in zip(detections['class_name'], detections.confidence)]

                

                bounding_box_annotator = sv.BoxAnnotator()
                label_annotator = sv.LabelAnnotator()

                annotated_image = bounding_box_annotator.annotate(scene=image, detections=detections)
                annotated_image = label_annotator.annotate(scene=annotated_image, detections=detections, labels=labels)

                annotated_image_filename = f"annotated_{filename}"
                annotated_image_path = os.path.join(app.config['UPLOAD_FOLDER'], annotated_image_filename)
                annotated_image.save(annotated_image_path)

                image_url = url_for('uploaded_file', filename=annotated_image_filename)
                formatted_labels = labels

        return render_template('admin_test.html', username=username, image_url=image_url, labels=formatted_labels)
    else:
        return redirect(url_for('login'))





@app.route('/get_disease_data', methods=['GET'])
def get_disease_data():
    if 'id' in session:
        try:
            user_id = session['id']
            filter_type = request.args.get('filter')

            if not filter_type:
                return jsonify({"error": "Filter type is required"}), 400

            today = datetime.datetime.today()

            # Special handling for "daily" filter
            if filter_type == 'daily':
                start_date = today - datetime.timedelta(days=1)

                query = """
                    SELECT hs.status_name, COUNT(*) AS disease_count
                    FROM detection d
                    INNER JOIN history h ON h.detection_id = d.detection_id
                    INNER JOIN health_status hs on d.health_status_id = hs.health_status_id
                    WHERE h.date_recorded >= %s AND d.user_id = %s AND NOT hs.status_name IN ('None Detected')
                    GROUP BY hs.status_name
                    ORDER BY disease_count DESC
                """
                cursor.execute(query, (start_date, user_id))
                results = cursor.fetchall()
                if not results:
                    return jsonify({"labels": [], "data": []})

                data = {
                    "labels": [result[0] for result in results],  # Disease names
                    "data": [result[1] for result in results]     # Disease counts
                }

                return jsonify(data)

            if filter_type == 'weekly':
                start_date = today - datetime.timedelta(weeks=1)
            elif filter_type == 'monthly':
                start_date = today - datetime.timedelta(days=30)
            elif filter_type == 'yearly':
                start_date = today - datetime.timedelta(days=365)
            else:
                return jsonify({"error": "Invalid filter type"}), 400

            query = """
                SELECT h.date_recorded, COUNT(*) AS detection_count
                FROM history h
                INNER JOIN detection d ON d.detection_id = h.detection_id
                WHERE h.date_recorded >= %s AND d.user_id = %s
                GROUP BY h.date_recorded
                ORDER BY h.date_recorded ASC
            """
            cursor.execute(query, (start_date, user_id))
            results = cursor.fetchall()


            if not results:
                return jsonify({"labels": [], "data": []})

            data = {
                "labels": [result[0].strftime('%Y-%m-%d') for result in results],  
                "data": [result[1] for result in results]                       
            }

            return jsonify(data)

        except Exception as e:
            print(f"Error in /get_disease_data: {e}")
            return jsonify({"error": "An error occurred on the server"}), 500



@app.route('/get_disease_data_farm', methods=['GET'])
def get_disease_data_farm():
    if 'id' in session:
        try:
            user_id = session['id']
            farm_id = request.args.get('farm_id')


            print(f"Farm ID: {farm_id}")

            today = datetime.datetime.today()
            start_date = today - datetime.timedelta(weeks=4) 

            print(f"Start Date: {start_date}")



            query = """
                SELECT h.date_recorded, COUNT(*) AS detection_count
                FROM history h
                INNER JOIN detection d ON d.detection_id = h.detection_id
                WHERE h.date_recorded >= %s AND d.farm_id = %s
                GROUP BY h.date_recorded
                ORDER BY h.date_recorded ASC
            """
            print(f"Executing query with params: start_date={start_date}, farm_id={farm_id}")
            
            cursor.execute(query, (start_date, farm_id))
            results = cursor.fetchall()

            print(f"Query Results: {results}")



            if not results:
                return jsonify({"labels": [], "data": []}), 404

            data = {
                "labels": [result[0].strftime('%Y-%m-%d') for result in results],
                "data": [result[1] for result in results]
            }

            return jsonify(data)

        except Exception as e:
            print(f"Error in /get_disease_data_farm: {e}")
            return jsonify({"error": "An error occurred on the server"}), 500
    else:
        return redirect(url_for('login'))







@app.route('/get_disease_stats', methods=['GET'])
def get_disease_stats():
    if 'id' in session:
        user_id = session['id']
        try:
            query = """
                SELECT hs.status_name, COUNT(*) AS detection_count
                FROM detection d
                INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                WHERE NOT hs.status_name IN ('None Detected') AND d.user_id = %s
                GROUP BY hs.status_name
                ORDER BY detection_count DESC
            """
            cursor.execute(query, (user_id,))
            results = cursor.fetchall()

            print("Results in stats:", results)

            if not results:
                return jsonify({"error": "No disease data found"}), 404

            data = {
                "labels": [result[0] for result in results],  # Disease names
                "data": [result[1] for result in results]  # Detection counts
            }

            print('Data in stats:', data)
            return jsonify(data)    
        
        except Exception as e:
            print(f"Error in /get_disease_stats: {e}")
            return jsonify({"error": "An error occurred on the server"}), 500
    else:
        return jsonify({"error": "Unauthorized access"}), 401

def compute_average_temperature(month):
    query = f"""
    SELECT 
        AVG(CAST(SUBSTRING_INDEX(h.weather_temp, '°C', 1) AS DECIMAL(10, 5))) AS average_celsius_temp,
        COUNT(d.health_status_id) AS detection_count
    FROM 
        detection d
    INNER JOIN 
        history h ON h.detection_id = d.detection_id
    WHERE 
        d.farm_id IS NOT NULL 
        AND DATE_FORMAT(h.date_recorded, '%Y-%m') = '{month}'  -- Filter by selected month
    """
    
    cursor.execute(query)
    result = cursor.fetchone()
    
    if result is None:
        return {
            'average_temperature': 0,
            'detection_count': 0
        }
    
    return {
        'average_temperature': result[0],
        'detection_count': result[1]
    }

@app.route('/weather', methods=['GET', 'POST'])
def weather():
    if request.method == 'POST':
        month = request.form.get('month') 
        weather_data = compute_average_temperature(month)

        print('Weather Data:', weather_data)

        return render_template('weather.html', weather_data=weather_data)

    return render_template('weather.html')

















@app.route('/dashboard_cp', methods=['GET', 'POST'])
def dashboard_cp():
    if 'id' in session:
        sid = session['id']
        farm_id = session.get('farm_id', None)
        print("farm_id = ", farm_id)

        cursor.execute("SELECT username FROM tbl_user WHERE user_id = %s", (sid,))
        result = cursor.fetchone()
        if result:
            username = result[0]
        else:
            return redirect(url_for('login'))

        temp_kelv = response['main']['temp']
        temp_cs, temp_fr = kelvin_to_cs_fr(temp_kelv)
        temp_cs = int(temp_cs)
        temp_fr = int(temp_fr) 

        weather = response['weather'][0]['description']
        icon = response['weather'][0]['icon']
        icon_url = f"http://openweathermap.org/img/wn/{icon}@2x.png"

        health_status = ''
        image_url = None
        formatted_labels = []
        results = []

        cursor.execute('SELECT u.user_id, u.role_id, tr.role_name FROM tbl_user u INNER JOIN tbl_role tr ON u.role_id = tr.role_id WHERE u.user_id = %s', (sid,))
        result = cursor.fetchone()
        role_name = result[2] if result else None
        print('this is role name', role_name)

        if check_person_in_ud_farm(sid):
            valid_rooms = 1
        else:
            valid_rooms = 0


        if role_name in ['user_farmer', 'farm_owner']:
            if farm_id is not None:  
                cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s AND uf.farm_id = %s  ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                            (sid, farm_id,))  

                result = cursor.fetchone()  
                status = result[0] if result else None  
                print("This is the status:", status)  
            else:
                cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                            (sid,))  

                result = cursor.fetchone()  
                status = result[0] if result else None  
                print("This is the status:", status) 
            
            if status == 'rejected':
                farm_id = None
                print(status)
                cursor.execute("""SELECT count(detection_id) FROM detection d
                              WHERE NOT d.health_status_id IN (8, 9) AND d.user_id = %s""", (sid,))
                ct = cursor.fetchone()

                cursor.execute("""SELECT count(detection_id), d.health_status_id, hs.status_name FROM detection d
                                INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                                WHERE NOT hs.status_name IN ('None Detected', 'healthy') AND d.user_id = %s
                                GROUP BY d.health_status_id
                                ORDER BY d.health_status_id DESC""", (sid,))
                max = cursor.fetchone()
                print("status is rejected")

            elif farm_id is None:
                farm_id = None
                print(status)
                cursor.execute("""SELECT count(detection_id) FROM detection d
                              WHERE NOT d.health_status_id IN (8, 9) AND d.user_id = %s""", (sid,))
                ct = cursor.fetchone()

                cursor.execute("""SELECT count(detection_id), d.health_status_id, hs.status_name FROM detection d
                                INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                                WHERE NOT hs.status_name IN ('None Detected', 'healthy') AND d.user_id = %s
                                GROUP BY d.health_status_id
                                ORDER BY d.health_status_id DESC""", (sid,))
                max = cursor.fetchone()
                print("status is rejected")

            else:
                cursor.execute("""SELECT count(detection_id) FROM detection d
                                INNER JOIN farm f ON d.farm_id = f.farm_id
                                WHERE NOT d.health_status_id IN (8, 9) AND f.farm_id = %s AND d.user_id = %s""", (farm_id, sid))
                ct = cursor.fetchone()

                cursor.execute("""SELECT count(detection_id), d.health_status_id, hs.status_name FROM detection d
                                INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                                WHERE NOT hs.status_name IN ('None Detected', 'healthy') AND d.farm_id = %s AND d.user_id = %s
                                GROUP BY d.health_status_id
                                ORDER BY d.health_status_id DESC""", (farm_id, sid))
                max = cursor.fetchone()
                print("status is owner or accepted")

                print('this is count max', ct, max)

        else:
            status = None
            cursor.execute("""SELECT count(detection_id) FROM detection d
                              WHERE NOT d.health_status_id IN (8, 9) AND d.user_id = %s""", (sid,))
            ct = cursor.fetchone()

            cursor.execute("""SELECT count(detection_id), d.health_status_id, hs.status_name FROM detection d
                              INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                              WHERE NOT d.health_status_id = 8 AND NOT hs.status_name = 'healthy' AND d.user_id = %s
                              GROUP BY d.health_status_id
                              ORDER BY d.health_status_id DESC""", (sid,))
            max = cursor.fetchone()
    
        max = max[2] if max else ''
        
        label = ''
        if request.method == 'POST':
            file = request.files['file']
            if file and allowed_file(file.filename):
                filename = file.filename
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                image = Image.open(file_path)
                

                
                model = get_model(model_id="calamansi-0bxbw/3")
                results = model.infer(image)
                detections = sv.Detections.from_inference(results[0].dict(by_alias=True, exclude_none=True))

                labels = [f"{class_name}" for class_name in zip(detections['class_name'])]

                bounding_box_annotator = sv.BoxAnnotator()
                label_annotator = sv.LabelAnnotator()

                annotated_image = bounding_box_annotator.annotate(scene=image, detections=detections)
                annotated_image = label_annotator.annotate(scene=annotated_image, detections=detections, labels=labels)

                sv.plot_image(annotated_image)

                annotated_image_filename = f"annotated_{filename}"
                annotated_image_path = os.path.join(app.config['UPLOAD_FOLDER'], annotated_image_filename)
                annotated_image.save(annotated_image_path)



                formatted_labels = [label.replace("('", "").replace("',)", "") for label in labels]

                image_url = url_for('uploaded_file', filename=annotated_image_filename)

                health_status_mapping = {
                    'black spot': 'black spots',
                    'greening': 'greening',
                    'scab': 'scab',
                    'thrips': 'thrips',
                    'heathy': 'healthy',
                    'None Detected': 'None Detected' 
                }

                print("Formatted labels: ", formatted_labels)
                processed_labels = set() 
                results = []  

                if formatted_labels == []:
                    label = 'None Detected'
                    print('label is ',label)
                    if label in health_status_mapping:
                        health_status_name = health_status_mapping[label]
                        print(f'Health Status name: {health_status_name}')
                        if status in ['rejected', 'member pending']:
                            farm_id = None

                        cursor.execute('SELECT cause, solution, description, health_status_id FROM health_status WHERE status_name = %s', (health_status_name,))
                        health_status = cursor.fetchone()
                        cursor.execute("INSERT INTO detection (user_id, img, prediction, health_status_id, farm_id) VALUES (%s, %s, %s, %s, %s)", 
                                (sid, image_url, label, health_status[3], farm_id))
                        connection.commit()
                        cursor.execute('SELECT detection_id FROM detection ORDER BY detection_id DESC LIMIT 1')
                        detection_id = cursor.fetchone()[0]

                        

                        time_recorded = datetime.datetime.now().strftime('%H:%M:%S')
                        date_recorded = datetime.datetime.now().strftime('%Y-%m-%d')
                        temperature = f"{temp_cs}°C / {temp_fr}°F"

                        cursor.execute('INSERT INTO history (date_recorded, time_recorded, detection_id, weather_temp, weather_status) VALUES (%s, %s, %s, %s, %s)',
                                            (date_recorded, time_recorded, detection_id, temperature, weather))
                        connection.commit()

                        if label not in processed_labels:
                            result = {
                                        'label': label,
                                        'description': health_status[2],
                                        'cause': health_status[0],
                                        'solution': health_status[1]
                                    }

                            results.append(result)

                            processed_labels.add(label)
                        
                else:
                    label = formatted_labels[0]


                    for label in formatted_labels:
                        if label in health_status_mapping:
                            health_status_name = health_status_mapping[label]
                            print(f'Health Status name: {health_status_name}')

                            if label == 'heathy':
                                label = 'healthy'

                            cursor.execute('SELECT cause, solution, description, health_status_id, source FROM health_status WHERE status_name = %s', (health_status_name,))
                            health_status = cursor.fetchone()

                            print('health status', health_status)
                            if health_status:
                                print(f"Health Status fetched: {health_status}")

                                if status in ['rejected', 'member pending']:
                                    farm_id = None

                                cursor.execute("INSERT INTO detection (user_id, img, prediction, health_status_id, farm_id) VALUES (%s, %s, %s, %s, %s)", 
                                (sid, image_url, label, health_status[3], farm_id))

                                connection.commit()

                                cursor.execute('SELECT detection_id FROM detection ORDER BY detection_id DESC LIMIT 1')
                                detection_id = cursor.fetchone()[0]

                                time_recorded = datetime.datetime.now().strftime('%H:%M:%S')
                                date_recorded = datetime.datetime.now().strftime('%Y-%m-%d')
                                temperature = f"{temp_cs}°C / {temp_fr}°F"

                                cursor.execute('INSERT INTO history (date_recorded, time_recorded, detection_id, weather_temp, weather_status) VALUES (%s, %s, %s, %s, %s)',
                                            (date_recorded, time_recorded, detection_id, temperature, weather))
                                connection.commit()

                                if label not in processed_labels:
                                    result = {
                                        'label': label,
                                        'description': health_status[2],
                                        'cause': health_status[0],
                                        'solution': health_status[1],
                                        'source':health_status[4]
                                    }

                                    results.append(result)
                                    processed_labels.add(label)

                                    print('THE RESULTS: ', results)
                            else:
                                print(f"No health status found for {health_status_name}")


        
        return render_template('dashboard_cp.html', username=username,count = ct[0], status = status,
                               health_status = health_status, 
                               temp_cs=int(temp_cs), temp_fr=int(temp_fr), weather=weather, icon=icon, icon_url=icon_url, 
                               image_url=image_url, formatted_labels=label, max = max,
                               members=members, user_type = role_name, results = results, valid_rooms= valid_rooms )
    else:
        return redirect(url_for('login'))





def add_trusted_location(path, office_version="16.0"):
    registry_path = f"Software\\Microsoft\\Office\\{office_version}\\Word\\Security\\Trusted Locations"

    try:
        key = reg.OpenKey(reg.HKEY_CURRENT_USER, registry_path, 0, reg.KEY_ALL_ACCESS)
        
        location_index = 0
        while True:
            try:
                reg.EnumKey(key, location_index)
                location_index += 1
            except OSError:
                break

        new_location_key = f"Location{location_index}"
        location_key = reg.CreateKey(key, new_location_key)

        reg.SetValueEx(location_key, "Path", 0, reg.REG_SZ, path)
        reg.SetValueEx(location_key, "AllowSubFolders", 0, reg.REG_DWORD, 1)
        reg.SetValueEx(location_key, "Description", 0, reg.REG_SZ, "Trusted Location added by Python script")

        reg.CloseKey(location_key)
        reg.CloseKey(key)

        print(f"Trusted location added: {path}")
    except Exception as e:
        print(f"Failed to add trusted location: {e}")

folder_path = r"C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output"
if os.path.exists(folder_path):
    add_trusted_location(folder_path)
else:
    print(f"Folder path does not exist: {folder_path}")







@app.route('/print_chart', methods=['POST'])
def print_chart():
    data = request.get_json()
    line_chart_data = data.get('line_chart')
    pie_chart_data = data.get('pie_chart')
    count = data.get('count')
    maxs = data.get('max')
    print("count max",count, maxs)
    farm_id = session['farm_id']


    today = datetime.datetime.today()
    cursor.execute('''
        SELECT h.weather_status, COUNT(d.detection_id) AS detection_count
        FROM detection d
        INNER JOIN history h ON h.detection_id = d.detection_id
        WHERE d.farm_id = %s
        GROUP BY h.weather_status
        ORDER BY detection_count DESC
        LIMIT 1
    ''',(farm_id,))
    weather_with_highest_detection = cursor.fetchone()


    if line_chart_data and pie_chart_data:
        line_chart_data = line_chart_data.split(',')[1]
        line_chart_bytes = base64.b64decode(line_chart_data)
        line_chart_image = Image.open(BytesIO(line_chart_bytes))
        line_chart_filename = 'line_chart_output.png'
        line_chart_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', line_chart_filename)
        line_chart_image.save(line_chart_path)

        pie_chart_data = pie_chart_data.split(',')[1]
        pie_chart_bytes = base64.b64decode(pie_chart_data)
        pie_chart_image = Image.open(BytesIO(pie_chart_bytes))
        pie_chart_filename = 'pie_chart_output.png'
        pie_chart_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', pie_chart_filename)
        pie_chart_image.save(pie_chart_path)

        doc = Document()
        
        doc.add_heading('Disease Detection Report', level=1)
        doc.add_paragraph(f"Date: {datetime.datetime.today().strftime('%B %d, %Y')}")  # Adding today's date in formal format
        
        doc.add_picture(line_chart_path, width=Inches(5))
        doc.add_paragraph("Figure 1: Analysis of disease detections over time displayed as a line chart. This chart provides a visual overview of disease occurrences recorded within the specified period.")

        doc.add_picture(pie_chart_path, width=Inches(5))
        doc.add_paragraph("Figure 2: Distribution of disease detections by category in a pie chart format. This figure highlights the proportional representation of each disease category observed.")

        doc.add_paragraph(
            f"The total number of disease occurrences detected within the specified period is {count}. "
            f"This metric reflects the extent of disease activity monitored at the farm and helps in assessing the prevalence of infections."
        )

        doc.add_paragraph(
            f"The most frequently detected disease during this period is '{maxs}'. This finding indicates a particular vulnerability or trend related to this disease within the farm's environment, suggesting that further preventive or remedial actions may be warranted."
        )
        
        if weather_with_highest_detection:
            weather_status, detection_count = weather_with_highest_detection
            doc.add_paragraph(
                f"The weather condition with the highest recorded disease detections is '{weather_status}', "
                f"with a total of  {detection_count} cases. This association between weather conditions and disease prevalence suggests that specific climatic factors may contribute to increased disease risk, which could inform future farming and prevention strategies."
            )
        else:
            doc.add_paragraph("No specific weather condition data was available for the recorded detections.")

        output_doc_filename = 'detection_report.docx'
        output_doc_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', output_doc_filename)
        doc.save(output_doc_path)

        return send_file(output_doc_path, as_attachment=True, download_name='report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    else:
        return jsonify({"error": "No chart data received"}), 400




@app.route('/check_status', methods=['GET'])
def check_status():
    if 'id' in session:
        sid = session['id']

        cursor.execute('''
            SELECT status FROM ud_farm 
            WHERE user_details_id = (SELECT user_details_id FROM user_details WHERE user_id = %s)
        ''', (sid,))
        current_status = cursor.fetchone()
        print('this is current status', current_status)

        if current_status:
            return jsonify({'status': current_status[0]})  
        else:
            return jsonify({'status': 'unknown'})
    else:
        return jsonify({'status': 'not_logged_in'}), 401
    





@app.route('/check_new_members', methods=['GET'])
def check_new_members():
    if 'farm_id' in session:
        farm_id = session['farm_id']
        if session['role_id'] == 2:

            cursor.execute('''
                SELECT u.username, ud.fname, ud.lname, uf.status, u.created_at
                FROM tbl_user u
                INNER JOIN user_details ud ON u.user_id = ud.user_id
                INNER JOIN ud_farm uf ON uf.user_details_id = ud.user_details_id
                WHERE uf.farm_id = %s
                ORDER BY u.created_at DESC LIMIT 1
            ''', (farm_id,))
            new_member = cursor.fetchone()
            print('this is new member', new_member)

            if new_member:
                member_name = f"{new_member[1]} {new_member[2]}"
                return jsonify({'new_member': member_name, 'status': new_member[3]})
            else:
                return jsonify({'new_member': None})
        else:
            return jsonify({'error': 'Not authorized'}), 401
            

    return jsonify({'error': 'Not authorized'}), 401





@app.route('/history', methods=['GET', 'POST'])
def history():
    if 'id' in session:
        sid = session['id']
        if 'farm_id' in session:
            farm_id = session['farm_id']
        else:
            farm_id = None

        cursor.execute("SELECT username FROM tbl_user WHERE user_id = %s", (sid,))
        username_result = cursor.fetchone()
        username = username_result[0] if username_result else "Unknown User"

        cursor.execute('''SELECT u.user_id, u.role_id, tr.role_id, tr.role_name 
                          FROM tbl_user u 
                          INNER JOIN tbl_role tr ON u.role_id = tr.role_id   
                          WHERE u.user_id = %s''', (sid,))
        role_result = cursor.fetchone()
        role_name = role_result[3] if role_result else None

        if farm_id is not None:  
            cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s AND uf.farm_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                            (sid, farm_id,))  

            result = cursor.fetchone()  
            status = result[0] if result else None  
            print("This is the status:", status)  
        else:
            cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                            (sid,))  

            result = cursor.fetchone()  
            status = result[0] if result else None  
            print("This is the status:", status)         


        if check_person_in_ud_farm(sid):
            valid_rooms = 1
             
        else:
            valid_rooms = 0

             
        farm_id = None
        if role_name in ['user_farmer', 'farm_owner']:
            cursor.execute('''SELECT u.user_id, u.username, ud.user_details_id, ud.fname, ud.lname, ud.mname, 
                              uf.farm_id, uf.status
                              FROM tbl_user u 
                              INNER JOIN user_details ud ON u.user_id = ud.user_id 
                              INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                              WHERE u.user_id = %s''', (sid,))
            result = cursor.fetchone()
            if result:
                farm_id = result[6]

        limit = " LIMIT 10 "
        see_all = False 
        close = " LIMIT 10 "

        custom_start_date = request.form.get('start_date')
        custom_end_date = request.form.get('end_date')


        if request.method == 'POST':
            see_all = request.form.get('see_all')
            close = request.form.get('close')

            if close:
                close = " LIMIT 10 "

            if custom_start_date and custom_end_date:
                query = """SELECT d.img, d.prediction, h.weather_status, h.weather_temp,
                                  h.date_recorded, h.time_recorded, hs.solution
                           FROM history h 
                           INNER JOIN detection d ON d.detection_id = h.detection_id 
                           INNER JOIN health_status hs ON hs.health_status_id = d.health_status_id
                           INNER JOIN tbl_user u ON u.user_id = d.user_id 
                           WHERE u.user_id = %s 
                             AND h.date_recorded BETWEEN %s AND %s
                           ORDER BY DATE(h.date_recorded), TIME(h.time_recorded) DESC"""
                params = (sid, custom_start_date, custom_end_date)
            elif see_all or close:
                query = """SELECT d.img, d.prediction, h.weather_status, h.weather_temp,
                                  h.date_recorded, h.time_recorded, hs.solution
                           FROM history h 
                           INNER JOIN detection d ON d.detection_id = h.detection_id 
                           INNER JOIN health_status hs ON hs.health_status_id = d.health_status_id
                           INNER JOIN tbl_user u ON u.user_id = d.user_id 
                           WHERE u.user_id = %s 
                           ORDER BY d.detection_id DESC""" + ('' if see_all else close)
                params = (sid,)

            else:
                query = """SELECT d.img, d.prediction, h.weather_status, h.weather_temp,
                                  h.date_recorded, h.time_recorded, hs.solution
                           FROM history h 
                           INNER JOIN detection d ON d.detection_id = h.detection_id 
                           INNER JOIN health_status hs ON hs.health_status_id = d.health_status_id
                           INNER JOIN tbl_user u ON u.user_id = d.user_id 
                           WHERE u.user_id = %s 
                           ORDER BY d.detection_id DESC"""
                params = (sid,)                

        else:
            query = """SELECT d.img, d.prediction, h.weather_status, h.weather_temp,
                              h.date_recorded, h.time_recorded, hs.solution
                       FROM history h 
                       INNER JOIN detection d ON d.detection_id = h.detection_id 
                       INNER JOIN health_status hs ON hs.health_status_id = d.health_status_id
                       INNER JOIN tbl_user u ON u.user_id = d.user_id 
                       WHERE u.user_id = %s 
                       ORDER BY d.detection_id DESC""" + limit
            params = (sid,)

        cursor.execute(query, params)
        h_records = cursor.fetchall()

        history = []
        for record in h_records:
            history.append({
                'image': record[0],
                'prediction': record[1],
                'weather_status': record[2],
                'weather_temp': record[3],
                'date_recorded': record[4],
                'time_recorded': record[5],
                'solution': record[6]
            })

        print("History:", history)
        history_message = "No history records found for the selected date range." if not history else None
        print(custom_start_date, custom_end_date)


        return render_template('history.html', user_id=sid, username=username, history=history, user_type=role_name, status=status, history_message=history_message, valid_rooms = valid_rooms)

    else:
        return redirect(url_for('login'))


@app.route('/filter_health_status', methods=['GET', 'POST'])
def filter_health_status():
    if 'id' in session:
        sid = session['id']



        if check_person_in_ud_farm(sid):
            valid_rooms = 1
             
        else:
            valid_rooms = 0


        cursor.execute("SELECT username FROM tbl_user WHERE user_id = %s", (sid,))
        username_result = cursor.fetchone()
        username = username_result[0] if username_result else "Unknown User"

        health_status = request.form.get('health_status')

        query = """SELECT d.img, d.prediction, h.weather_status, h.weather_temp,
                          h.date_recorded, h.time_recorded, hs.solution
                   FROM history h 
                   INNER JOIN detection d ON d.detection_id = h.detection_id 
                   INNER JOIN health_status hs ON hs.health_status_id = d.health_status_id
                   INNER JOIN tbl_user u ON u.user_id = d.user_id 
                   WHERE u.user_id = %s"""
        params = [sid]

        if health_status:
            query += " AND hs.status_name = %s"
            params.append(health_status)

        query += " ORDER BY h.date_recorded DESC, h.time_recorded DESC"

        cursor.execute(query, tuple(params))
        h_records = cursor.fetchall()

        history = [
            {
                'image': record[0],
                'prediction': record[1],
                'weather_status': record[2],
                'weather_temp': record[3],
                'date_recorded': record[4],
                'time_recorded': record[5],
                'solution': record[6],
            }
            for record in h_records
        ]

        history_message = (
            "No history records found for the selected health status."
            if not history
            else None
        )

        return render_template(
            'history.html',
            user_id=sid,
            username=username,
            history=history,
            history_message=history_message,
            valid_rooms = valid_rooms
        )

    return redirect(url_for('login'))


    
@app.route('/add_farm', methods=['GET', 'POST'])
def add_farm():
    if 'id' in session:
        user_id = session['id']
        farm_id = session['farm_id']

        text = ''
        if farm_id is not None:  
            cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s AND uf.farm_id = %s ORDER BY uf.ud_farm_id DESC LIMIT 1''', 
                            (user_id, farm_id,))  

            result = cursor.fetchone()  
            status = result[0] if result else None  
            print("This is the status:", status)  
        else:
            cursor.execute('''SELECT uf.status 
                                FROM tbl_user u
                                INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                WHERE u.user_id = %s uf.ud_farm_id ORDER BY uf.ud_farm_id LIMIT 1''', 
                            (user_id,))  

            result = cursor.fetchone()  
            status = result[0] if result else None  
            print("This is the status:", status)    
        if check_person_in_ud_farm(user_id):
            valid_rooms = 1
             
        else:
            valid_rooms = 0

             
        cursor.execute('''SELECT ud.user_details_id, uf.status FROM user_details ud 
                           inner join tbl_user u on u.user_id = ud.user_id
                           inner join ud_farm uf on uf.user_details_id = ud.user_details_id
                            WHERE u.user_id = %s''', (user_id,))
        res = cursor.fetchone()
        if not res:
            cursor.execute('''SELECT ud.user_details_id FROM user_details ud  
                           inner join tbl_user u on u.user_id = ud.user_id
                            WHERE u.user_id = %s''', (user_id,))
            res = cursor.fetchone()
            user_details_id = res[0]
        else:
            user_details_id = res[0]


        cursor.execute("SELECT r.role_name FROM tbl_role r INNER JOIN tbl_user u on u.role_id = r.role_id WHERE u.user_id = %s", (user_id,))
        role_name = cursor.fetchone()[0]

        if request.method == 'POST':

            farm_name = request.form['farm_name']
            province = request.form['province-dropdown']
            city = request.form['city-dropdown']
            brgy = request.form['barangay-dropdown']
            strt_add = request.form['strt_add']
            zip_code = request.form['zip-code']
            f_img = request.files['f_img']


            if f_img:
                filename = f_img.filename
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                f_img.save(file_path)
                image_url = url_for('uploaded_file', filename=filename)

            else:
                image_url = None



            country = 'Philippines'
            address = f"{brgy}, {city}, {province}, {country}"
            print(f"Geocoding address: {address}")

            geolocator = Nominatim(user_agent="JOSHAPP")
            location = geolocator.geocode(address)
            print(location)

            if location:
                latitude = location.latitude
                longitude = location.longitude
                cursor.execute('INSERT INTO farm (farm_name, province, city, brgy, strt_add, zip_code, latitude, longitude, f_img, farm_status) VALUES (%s, %s, %s, %s, %s, %s, %s,%s, %s, "active")', 
                        (farm_name, province, city, brgy, strt_add, zip_code, latitude, longitude, image_url))
                connection.commit()
                cursor.execute('SELECT farm_id FROM farm ORDER BY farm_id DESC LIMIT 1')
                farm_id = cursor.fetchone()[0]
                group_code = generate_random_code()
                cursor.execute('INSERT INTO tbl_group (group_code) VALUES( %s)', (group_code,))
                connection.commit()
                cursor.execute('SELECT g.group_id FROM tbl_group g ORDER BY g.group_id DESC LIMIT 1')
                group = cursor.fetchone()
                group_id = group[0]
                print('new group code', group_code)
                cursor.execute('INSERT INTO ud_farm VALUES(NULL, %s, %s, "Owner", %s)', (user_details_id, farm_id, group_id,))
                connection.commit()
                cursor.execute('''UPDATE tbl_user  SET role_id = 6 WHERE user_id = %s''',(user_id,))
                connection.commit()

                
                text= 'Farm Added'
            else:
                text = f'No Location {address}'

        return render_template('add_farm.html', user_id = user_id, text=text, user_type = role_name, valid_rooms = valid_rooms,  status=status)
    else:
        return redirect(url_for('login'))
        

def check_person_in_ud_farm(user_id):
    cursor.execute('''
        SELECT COUNT(*)
        FROM ud_farm uf inner join user_details ud on uf.user_details_id = ud.user_details_id
        inner join tbl_user u on u.user_id = ud.user_id
        WHERE u.user_id = %s
    ''', (user_id,))

    result = cursor.fetchone()

    if result[0] > 0:
        return True
    else:
        return False



@app.route('/rooms', methods=['GET', 'POST'])
def rooms():
    if 'id' in session:
        user_id = session['id']
        username = session['username']

        cursor.execute('''SELECT g.group_code, f.farm_name
                          FROM tbl_user u
                          INNER JOIN user_details ud ON u.user_id = ud.user_id
                          INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id
                       inner join farm f on f.farm_id = uf.farm_id
                          INNER JOIN tbl_group g ON uf.group_id = g.group_id
                          WHERE u.user_id = %s AND uf.status IN ('Owner', 'member')''', (user_id,))
        rooms = cursor.fetchall()

        cursor.execute('''SELECT tr.role_name, uf.status FROM tbl_role tr
                       inner join tbl_user u on u.role_id = tr.role_id
                       inner join user_details ud on ud.user_id = u.user_id
                       inner join ud_farm uf on uf.user_details_id = ud.user_details_id
                       WHERE u.user_id = %s''',(user_id,))
        ress = cursor.fetchone()
        role_name = ress[0]
        status = ress[1]

        selected_room = None
        sanitized_messages = []
        ud_farm_id = None


        if check_person_in_ud_farm(user_id):
            valid_rooms = 1
             
        else:
            valid_rooms = 0

             



        if request.method == 'POST':
            selected_room = request.form.get('room')
            if selected_room:
                cursor.execute('''SELECT u.username, m.message, m.timestamp
                                  FROM tbl_user u 
                                  INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                  INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                  INNER JOIN tbl_message m ON m.ud_farm_id = uf.ud_farm_id
                                  INNER JOIN tbl_group g ON uf.group_id = g.group_id
                                  WHERE g.group_code = %s''', (selected_room,))
                messages = cursor.fetchall()

                sanitized_messages = []
                for msg in messages:
                    print('message is :', msg)
                    username, message, timestamp = msg
                    if contains_inappropriate_content(message):
                        sanitized_message = ("This message has been flagged for inappropriate content.",)
                    else:
                        sanitized_message = (message,)
                    
                    sanitized_messages.append((username, sanitized_message[0], timestamp))
                    print(sanitized_messages)


                cursor.execute('''SELECT u.username, ud.fname, ud.lname, uf.farm_id, ud.user_details_id
                                  FROM tbl_user u 
                                  INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                  INNER JOIN ud_farm uf ON uf.user_details_id = ud.user_details_id
                                  WHERE u.user_id = %s''', (user_id,))
                user = cursor.fetchone()

                cursor.execute('''SELECT uf.ud_farm_id ,g.group_id
                                  FROM tbl_group g
                                  INNER JOIN ud_farm uf ON uf.group_id = g.group_id
                                  INNER JOIN user_details ud ON uf.user_details_id = ud.user_details_id
                                  WHERE g.group_code = %s AND ud.user_details_id = %s ''', 
                               (selected_room, user[4]))
                res = cursor.fetchone()
                print(res)
                if res:
                    ud_farm_id = res[0]

        return render_template('rooms.html', rooms=rooms, selected_room=selected_room, messages=sanitized_messages, username=username, ud_farm_id=ud_farm_id
                               ,user_type = role_name,status = status, valid_rooms = valid_rooms)

    return redirect(url_for('login'))





@app.route('/chat/<room>')
def chat(room):
    if 'id' in session:
        user_id = session['id']
        
        cursor.execute('''SELECT u.username, ud.fname, ud.lname, uf.farm_id, ud.user_details_id
                          FROM tbl_user u 
                          INNER JOIN user_details ud ON u.user_id = ud.user_id 
                        inner join ud_farm uf on uf.user_details_id = ud.user_details_id
                          WHERE u.user_id = %s''', (user_id,))
        user = cursor.fetchone()
        print(user)

        cursor.execute('''SELECT uf.ud_farm_id ,g.group_id
                              FROM tbl_group g
                       inner join ud_farm uf on uf.group_id = g.group_id
                       inner join user_details ud on uf.user_details_id = ud.user_details_id
                              WHERE g.group_code = %s and ud.user_details_id = %s ''', (room, user[4], ))
        res = cursor.fetchone()
        ud_farm_id = res[0]



        
        
        if user:
            username = user[0]
            full_name = f"{user[1]} {user[2]}"
            
            cursor.execute('''SELECT u.username, m.message, m.timestamp
                                    FROM tbl_user u 
                                    INNER JOIN user_details ud ON u.user_id = ud.user_id 
                                    INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                                    INNER JOIN tbl_message m ON m.ud_farm_id = uf.ud_farm_id
                                    INNER JOIN tbl_group g ON uf.group_id = g.group_id
                                    WHERE g.group_code = %s
                              ''', (room,))
            messages = cursor.fetchall()
            


            return render_template('chat.html', username=username, full_name=full_name, room=room, ud_farm_id=ud_farm_id, messages=messages)
    
    return redirect(url_for('login'))



@socketio.on('join')
def on_join(data):
    username = data['username']
    room = data['room']
    join_room(room)
    send({"username": username, "msg": "has entered the room", "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}, to=room)

@socketio.on('leave')
def on_leave(data):
    username = data['username']
    room = data['room']
    leave_room(room)
    send({"username": username, "msg": "has left the room", "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}, to=room)


inappropriate_words = ['tite', 'burat', 'puke','pokpok','pepe', 'bilat', 'putangina']

def contains_inappropriate_content(message):
    for word in inappropriate_words:
        if word.lower() in message.lower():
            return True
    return False

@socketio.on('message')
def handle_message(data):
    message = data['msg']
    room = data['room']
    ud_farm_id = data['ud_farm_id']
    data['timestamp'] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if contains_inappropriate_content(message):
        data['msg'] = "This message has been flagged for inappropriate content."
        message = data['msg']

    cursor.execute('INSERT INTO tbl_message (message, timestamp, ud_farm_id) VALUES (%s, %s, %s)', 
                   (message, data['timestamp'], ud_farm_id ))
    connection.commit()

    send({'msg': data['msg'], 'username': data['username'], 'timestamp': data['timestamp']}, to=room, broadcast=True)




@app.route('/farm', methods=['GET', 'POST'])
def farm():
    if 'id' not in session:
        return redirect(url_for('login'))
    
    sid = session['id']
    cursor.execute('''SELECT uf.status FROM tbl_user u
                              INNER JOIN user_details ud ON u.user_id = ud.user_id 
                              INNER JOIN ud_farm uf ON ud.user_details_id = uf.user_details_id 
                              WHERE u.user_id = %s ORDER BY uf.farm_id DESC LIMIT 1''', (sid,))
    result = cursor.fetchone()
    status = result[0] if result else None
    print("this is status", status)
    if status in ['member pending', 'rejected']:
        farm_id = None
        session['farm_id'] = farm_id
        return redirect(url_for('dashboard_cp'))




    sid = session['id']
    role_id = session.get('role_id', None)

    cursor.execute('''SELECT f.farm_id, f.farm_name 
                      FROM farm f 
                      INNER JOIN ud_farm uf ON uf.farm_id = f.farm_id
                      INNER JOIN user_details ud ON uf.user_details_id = ud.user_details_id 
                      WHERE ud.user_id = %s AND NOT uf.status = "rejected" AND NOT uf.status = "member pending" 
                   ''', (sid,))
    farms = cursor.fetchall()

    farms_with_counts = []

    if not farms:
        cursor.execute('''SELECT uf.farm_id 
                          FROM ud_farm uf 
                          INNER JOIN user_details ud ON uf.user_details_id = ud.user_details_id
                          INNER JOIN tbl_user u ON u.user_id = ud.user_id  
                          WHERE u.user_id = %s 
                          ORDER BY u.user_id DESC LIMIT 1''', (sid,))
        farm_id = cursor.fetchone()

        if not farm_id:
            return redirect(url_for('dashboard_cp'))
        else:
            session['farm_id'] = farm_id[0]
            return redirect(url_for('dashboard_cp'))

    for farm in farms:
        farm_id, farm_name = farm

        cursor.execute('''SELECT COUNT(d.detection_id) 
                          FROM detection d 
                          WHERE d.farm_id = %s AND d.user_id = %s''', (farm_id, sid))
        detection_count = cursor.fetchone()[0]

        farms_with_counts.append({
            'farm_id': farm_id,
            'farm_name': farm_name,
            'detection_count': detection_count
        })

    session['farm_id'] = farms_with_counts[0]['farm_id']
    print(f"Default farm_id: {session['farm_id']}")

    if request.method == 'POST':
        farm_id = request.form['farm_id']
        session['farm_id'] = farm_id
        print(f"Selected farm_id: {session['farm_id']}")
        return redirect(url_for('dashboard_cp'))

    return render_template('farm.html', farms=farms_with_counts, role_id=role_id)








@app.route('/leave_farm', methods=['POST'])
def leave_farm():
    if 'id' in session:
        sid = session['id']
        farm_id = request.form['farm_id']

        cursor.execute('SELECT role_id FROM tbl_user WHERE user_id = %s', (sid,))
        role_id = cursor.fetchone()[0]

        if role_id == 5:
            cursor.execute('''
                    DELETE m FROM tbl_message m INNER JOIN ud_farm uf ON m.ud_farm_id = uf.ud_farm_id
                    WHERE uf.farm_id = %s
                ''', (farm_id,))
            connection.commit()

            cursor.execute('''
                DELETE FROM ud_farm 
                WHERE user_details_id = (SELECT user_details_id FROM user_details WHERE user_id = %s) 
                AND farm_id = %s
            ''', (sid, farm_id))
            connection.commit()

            flash('You have successfully left the farm.', 'success')
        
        elif role_id == 6:
            reason = request.form['reason']

            if reason:
                
                cursor.execute('''
                    DELETE m FROM tbl_message m INNER JOIN ud_farm uf ON m.ud_farm_id = uf.ud_farm_id
                    WHERE uf.farm_id = %s
                ''', (farm_id,))
                connection.commit()

                cursor.execute('''
                    DELETE FROM ud_farm 
                    WHERE farm_id = %s
                ''', (farm_id,))
                connection.commit()


                cursor.execute('''UPDATE farm  SET farm_status ="Pending For Deletion" WHERE farm_id = %s''',(farm_id,))
                connection.commit()



                flash('You have successfully deleted the farm. All associated farmers have been removed.', 'success')
            else:
                flash('You must provide a reason to delete the farm.', 'danger')

        return redirect(url_for('farm'))
    else:
        return redirect(url_for('login'))




@app.route('/adminfarm', methods=['GET', 'POST'])
def adminfarm():
    if 'id' in session:
        sid = session['id']
        
        cursor.execute('''SELECT f.farm_id, f.farm_name, f.f_img FROM farm f''')
        farms = cursor.fetchall()

        farm_with_images = []
        h_index = 0 

        for f in farms:
            farm_id = f[0]
            farm_name = f[1]
            farm_img_binary = f[2] 

            if farm_img_binary:
                img = farm_img_binary
                print(img)

            else:
                img = None 

            farm_with_images.append((farm_id, farm_name, img))

        return render_template('adminfarm.html', farms=farm_with_images)

    else:
        return redirect(url_for('login'))











@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/')
def index():


    

    return render_template('index.html')



@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form:
        username = request.form['username']
        password = request.form['password']

        cursor.execute("SELECT * FROM tbl_user WHERE username = %s LIMIT 1", (username,))
        user = cursor.fetchone()

        if user and user[5] == password: 
            role_id = user[2]
            print('ang role_id', role_id)

            cursor.execute("SELECT * FROM user_details WHERE user_id = %s LIMIT 1", (user[0],))
            user_details = cursor.fetchone()

            if role_id in [5, 6]:
                cursor.execute("SELECT * FROM ud_farm WHERE user_details_id = %s LIMIT 1", (user_details[0],))
                check_ud_farm = cursor.fetchone()
            

            
            if role_id == 7:
                if user_details:
                    flash('Please continue your registration')
                    return redirect(url_for('ask', user_id=user[0]))
                else:
                    flash('Please continue your registration')
                    return redirect(url_for('farmback', user_id=user[0]))
            elif role_id == 6 and not check_ud_farm:
                flash('Please continue your registration')
                return redirect(url_for('register_farm', user_id=user[0]))
            
            else:
                session['loggedin'] = True
                session['id'] = user[0]
                session['username'] = user[1]
                session['role_id'] = role_id

                if role_id == 6:
                    


                    return redirect(url_for('farm'))
                elif role_id == 2:
                    session['admin'] = True
                    return redirect(url_for('admin_dashboard'))
                elif role_id == 1:
                    return redirect(url_for('dashboard_cp'))
                elif role_id == 5:
                    
                    return redirect(url_for('farm'))

        else:
            text = 'Incorrect username or password!'
            return render_template('login.html', show_modal=True, modal_message=text)
    else:
        return render_template('login.html', show_modal=False)



@app.route('/admin', methods=['GET'])
def admin_dashboard():
    if 'id' in session:
        sid = session['id']
        cursor.execute("SELECT role_id FROM tbl_user WHERE user_id = %s", (sid,))
        role_id = cursor.fetchone()[0]

        if role_id == 2: 
            cursor.execute('SELECT COUNT(user_id) FROM tbl_user')
            total_users = int(cursor.fetchone()[0])
            cursor.execute('SELECT COUNT(detection_id) FROM detection')
            total_detection = int(cursor.fetchone()[0])
            cursor.execute('SELECT COUNT(farm_id) FROM farm Where farm_status = "active" OR farm_status = "Active"')
            active_farms = int(cursor.fetchone()[0])
            cursor.execute('SELECT COUNT(farm_id) FROM farm Where farm_status = "Pending for Deletion"')
            pending_farms = int(cursor.fetchone()[0])
            cursor.execute('''SELECT u.username, ud.fname, ud.lname, ud.mname, r.role_name 
                              FROM tbl_user u 
                              JOIN tbl_role r ON u.role_id = r.role_id 
                              INNER JOIN user_details ud ON u.user_id = ud.user_id''')
            users = cursor.fetchall()

            cursor.execute('''SELECT COUNT(ud.user_details_id) AS member_count, f.farm_name, uf.status, f.province, f.city, f.brgy, f.strt_add 
                              FROM user_details ud
                              INNER JOIN ud_farm uf ON uf.ud_farm_id = ud.user_details_id
                              INNER JOIN farm f ON uf.farm_id = f.farm_id
                              GROUP BY f.farm_name''')
            farms = cursor.fetchall()

            return render_template('admin_dashboard.html', 
                                   users=users, 
                                   farms=farms, total_detection = total_detection,  active_farms = active_farms, pending_farms = pending_farms, total_users = total_users)
        else:
            return "Access denied", 403
    else:
        return redirect(url_for('login'))


@app.route('/print_chart_admin', methods=['POST'])
def print_chart_admin():
   
    data = request.get_json()
    line_chart_data = data.get('line_chart')
    pie_chart_data = data.get('pie_chart')
    performance_chart_data = data.get('performance_chart')


    if line_chart_data and pie_chart_data:
        line_chart_data = line_chart_data.split(',')[1] 
        line_chart_bytes = base64.b64decode(line_chart_data)

        pie_chart_data = pie_chart_data.split(',')[1]
        pie_chart_bytes = base64.b64decode(pie_chart_data)

        performance_chart_data = performance_chart_data.split(',')[1]
        performance_chart_bytes = base64.b64decode(performance_chart_data)

        line_chart_image = Image.open(BytesIO(line_chart_bytes))
        line_chart_filename = 'line_chart_output.png'
        line_chart_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', line_chart_filename)

        line_chart_image.save(line_chart_path)

        pie_chart_image = Image.open(BytesIO(pie_chart_bytes))
        pie_chart_filename = 'pie_chart_output.png'
        pie_chart_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', pie_chart_filename)
        pie_chart_image.save(pie_chart_path)


        performance_chart_image = Image.open(BytesIO(performance_chart_bytes))
        performance_chart_filename = 'performance_chart_output.png'
        performance_chart_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', performance_chart_filename)
        performance_chart_image.save(performance_chart_path)

        doc = Document()
        doc.add_heading('Disease Detection Report', level=1)
        
        doc.add_picture(line_chart_path, width=Inches(5)) 
        doc.add_picture(pie_chart_path, width=Inches(5))  
        doc.add_picture(performance_chart_path, width=Inches(5))  

        output_doc_filename = 'detection_report.docx'
        output_doc_path = os.path.join(r'C:\xampp\htdocs\backup capstone\xammp_projects\static\print_output', output_doc_filename)
        doc.save(output_doc_path)

        return send_file(output_doc_path, as_attachment=True, download_name='report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    else:
        return jsonify({"error": "No chart data received"}), 400



@app.route('/admin_statistics', methods=['GET', 'POST'])
def admin_statistics():
    if 'id' in session:
        sid = session['id']
        cursor.execute("SELECT role_id FROM tbl_user WHERE user_id = %s", (sid,))
        role_id = cursor.fetchone()[0]

        if role_id == 2:  
            period = request.args.get('period', 'monthly') 
            today = datetime.datetime.today()
            if period == 'weekly':
                start_date = today - datetime.timedelta(weeks=1)
            elif period == 'monthly':
                start_date = today - datetime.timedelta(weeks=4)
            elif period == 'yearly':
                start_date = today - datetime.timedelta(days=365)
            else:
                start_date = today - datetime.timedelta(weeks=4)  

            today = datetime.datetime.today()
            custom_start_date = request.args.get('start_date')
            custom_end_date = request.args.get('end_date')
            if custom_start_date and custom_end_date:
                try:
                    start_date = datetime.datetime.strptime(custom_start_date, '%Y-%m-%d')
                    end_date = datetime.datetime.strptime(custom_end_date, '%Y-%m-%d')
                except ValueError:
                    return jsonify({"error": "Invalid date format. Use YYYY-MM-DD."}), 400
            else:
                end_date = today

            query = """
                SELECT f.farm_name, COUNT(d.detection_id) AS detection_count
                FROM farm f
                LEFT JOIN detection d ON f.farm_id = d.farm_id
                LEFT JOIN history h ON h.detection_id = d.detection_id
                WHERE h.date_recorded >= %s AND h.date_recorded <= %s OR h.date_recorded IS NULL
                GROUP BY f.farm_name
                ORDER BY detection_count DESC

            """
            cursor.execute(query, (start_date, end_date))
            detection_results = cursor.fetchall()

            farm_labels = [row[0] for row in detection_results]
            detection_data = [row[1] for row in detection_results]

            query2 = """SELECT hs.status_name, COUNT(*) AS detection_count
                        FROM detection d
                        INNER JOIN health_status hs ON d.health_status_id = hs.health_status_id
                        LEFT JOIN history h ON h.detection_id = d.detection_id
                        WHERE hs.status_name NOT IN ('None Detected', 'healthy') AND h.date_recorded >= %s AND h.date_recorded <= %s OR h.date_recorded IS NULL
                        GROUP BY hs.status_name
                        ORDER BY detection_count DESC"""
            cursor.execute(query2, (start_date, end_date))
            results2 = cursor.fetchall()

            health_stat_labels = [row[0] for row in results2]
            health_stat_data = [row[1] for row in results2]

            cursor.execute('SELECT COUNT(user_id) FROM tbl_user WHERE role_id NOT IN (2)')
            total_users = cursor.fetchone()[0]

            cursor.execute('SELECT COUNT(user_id) FROM tbl_user WHERE role_id = 5') 
            total_farmer = int(cursor.fetchone()[0])
            cursor.execute('SELECT COUNT(user_id) FROM tbl_user WHERE role_id = 6') 
            total_farmer_owner = int(cursor.fetchone()[0])

            cursor.execute('SELECT COUNT(user_id) FROM tbl_user WHERE role_id = 1')
            total_backyard = int(cursor.fetchone()[0])

            


            user_data = {
                'labels': ['Farmers', 'Backyard'],
                'datasets': [{
                    'label': 'Number of Users',
                    'backgroundColor': ['rgba(255, 99, 132, 0.2)', 'rgba(54, 162, 235, 0.2)','rgba(60, 179, 113, 0.2)'],
                    'borderColor': ['rgba(255, 99, 132, 1)', 'rgba(54, 162, 235, 1)','rgba(60, 179, 113, 1)'],
                    'borderWidth': 1,
                    'data': [total_farmer, total_backyard, total_farmer_owner]
                }]
            }

            chart_data = {
                'labels': farm_labels,
                'datasets': [{
                    'label': 'Disease Detections',
                    'backgroundColor': 'rgba(54, 162, 235, 0.2)',
                    'borderColor': 'rgba(54, 162, 235, 1)',
                    'borderWidth': 1,
                    'data': detection_data
                }]
            }

            health_data = {
                'labels': health_stat_labels,
                'datasets': [{
                    'label': 'Disease Count',
                    'backgroundColor': 'rgba(255, 159, 64, 0.2)',
                    'borderColor': 'rgba(255, 159, 64, 1)',
                    'borderWidth': 1,
                    'data': health_stat_data
                }]
            }
            return render_template('admin_statistics.html', 
                                   chart_data=user_data, 
                                   farm_chart_data=chart_data,
                                   health_data=health_data,
                                   total_users=total_users,
                                   )
        else:
            return "Access denied", 403
    else:
        return redirect(url_for('login'))






        
@app.route('/admin_update', methods=['POST'])
def admin_update():
    if 'admin' in session:
        user_id = request.form['user_id']
        user_details_id = request.form['user_details_id']

        username = request.form.get('username', '')
        password = request.form.get('password', '')
        fname = request.form.get('fname', '')
        mname = request.form.get('mname', '')
        lname = request.form.get('lname', '')
        province = request.form.get('province', '')
        city = request.form.get('city', '')
        brgy = request.form.get('brgy', '')
        strt_add = request.form.get('strt_add', '')
        cont_num = request.form.get('cont_num', '')

        cursor.execute('''
            UPDATE tbl_user
            SET username = %s, password = %s
            WHERE user_id = %s
        ''', (username, password, user_id))

        cursor.execute('''
            UPDATE user_details
            SET fname = %s, mname = %s, lname = %s, province = %s, city = %s, brgy = %s, strt_add = %s, cont_num = %s
            WHERE user_details_id = %s
        ''', (fname, mname, lname, province, city, brgy, strt_add, cont_num, user_details_id))

        connection.commit()
        text = 'User updated successfully!'
        return redirect(url_for('admin_crud', text=text))
    else:
        return redirect(url_for('login'))





@app.route('/admin_delete', methods=['POST'])
def admin_delete():
    if 'admin' in session:
        user_id = request.form['user_id']
        user_details_id = request.form['user_details_id']


        cursor.execute('''SELECT uf.user_details_id FROM ud_farm uf INNER JOIN user_details ud on uf.user_details_id = ud.user_details_id
        INNER JOIN tbl_user u on u.user_id = ud.user_id WHERE u.user_id = %s''', (user_id,))
        check_farm = cursor.fetchone()

        print(check_farm)

        if check_farm:
            cursor.execute('''
                    DELETE m FROM tbl_message m INNER JOIN ud_farm uf ON m.ud_farm_id = uf.ud_farm_id
                    WHERE uf.user_details_id = %s
                ''', (user_details_id,))
            connection.commit()
            cursor.execute('DELETE FROM ud_farm WHERE user_details_id = %s', (user_details_id,))
            connection.commit()




        cursor.execute('DELETE FROM user_details WHERE user_id = %s', (user_id,))
        connection.commit()
        cursor.execute('DELETE FROM detection WHERE user_id = %s', (user_id,))
        connection.commit()
        cursor.execute('DELETE h FROM history h inner join detection d on d.detection_id = h.detection_id WHERE user_id = %s', (user_id,))
        connection.commit()
        cursor.execute('DELETE FROM tbl_user WHERE user_id = %s', (user_id,))
        connection.commit()
        text = ('User deleted successfully!')
        text = 'User deleted successfully!'
        return redirect(url_for('admin_crud', text=text))

    else:
        return redirect(url_for('login'))


@app.route('/admin_crud', methods=['GET', 'POST'])
def admin_crud():
    if 'admin' in session:

        cursor.execute('''SELECT u.username, u.password, ud.fname, ud.mname, ud.lname, ud.province, ud.city, ud.brgy, ud.strt_add, ud.cont_num, u.user_id, ud.user_details_id
        FROM tbl_user u INNER JOIN user_details ud ON ud.user_id = u.user_id
        ''')
        results = cursor.fetchall()

        text = request.args.get('text', None)

        return render_template('admin_crud.html', results = results, text = text)
    else:
        return redirect(url_for('login')) 























@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))




if __name__ == '__main__':
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)
